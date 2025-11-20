import hashlib
from BD_models import User
import requests
import base64
from io import BytesIO
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH


### Orders

# Хеширование пароля
def hash_password(password):
    return hashlib.sha256(password.encode('utf-8')).hexdigest()

# Проверка пользователя
def check_user(result):
    if result.cookies.get("login") and result.cookies.get("password"):
        login = result.cookies.get("login").encode('latin-1').decode('utf-8')
        password = result.cookies.get("password")
        user = User.query.filter_by(login=login, password=hash_password(password)).first()
        return user
    return None


### Kandidat

# Сохранение анкеты в docx
def save_to_docx(form_data):
    doc = Document()

    # Настройка стиля по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Заголовок
    heading = doc.add_heading('АНКЕТА СОИСКАТЕЛЯ', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Mm(10)

    # Основные данные
    doc.add_heading('Основные данные', level=1)

    basic_fields = [
        ("ФИО", form_data['personalInfo'].get("fullName", "Не указано")),
        ("Должность", form_data['personalInfo'].get("position", "Не указано")),
        ("Дата рождения", form_data['personalInfo'].get("birthDate", "Не указано")),
        ("Место рождения", form_data['personalInfo'].get("birthPlace", "Не указано")),
        ("Место жительства", form_data['personalInfo'].get("address", "Не указано")),
        ("Гражданство", form_data['personalInfo'].get("citizenship", "Не указано")),
        ("Телефон", form_data['personalInfo'].get("phone", "Не указано")),
        ("Email", form_data['personalInfo'].get("email", "Не указано")),
        ("Семейное положение", form_data['personalInfo'].get("familyStatus", "Не указано")),
        ("Паспорт серия", form_data['personalInfo'].get("passport_ser", "Не указано")),
        ("Паспорт номер", form_data['personalInfo'].get("passport_num", "Не указано")),
        ("Паспорт код подразделения", form_data['personalInfo'].get("passport_cod", "Не указано")),
        ("Паспорт дата", form_data['personalInfo'].get("passport_date", "Не указано")),
        ("Паспорт выдан", form_data['personalInfo'].get("passport_otdel", "Не указано")),
        ("Автомобиль", form_data['personalInfo'].get("autonumber", "Не указано")),
        ("Узнал о вакансии", form_data['personalInfo'].get("vacancySource", "Не указано")),
        ("Желаемая зарплата", form_data['personalInfo'].get("desiredSalary", "Не указано")),
    ]

    for label, value in basic_fields:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
        p.paragraph_format.space_after = Mm(3)

    # Образование
    doc.add_heading('Образование', level=1)
    education_data = form_data.get("education", [])
    if education_data:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Учебное заведение'
        hdr_cells[1].text = 'Начало'
        hdr_cells[2].text = 'Окончание'
        hdr_cells[3].text = 'Специальность'

        for item in education_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("institution", "—")
            row_cells[1].text = item.get("startDate", "—")
            row_cells[2].text = item.get("endDate", "—")
            row_cells[3].text = item.get("specialty", "—")
    else:
        doc.add_paragraph("Образование не указано")

    # Опыт работы
    doc.add_heading('Опыт работы', level=1)
    work_data = form_data.get("workExperience", [])
    if work_data:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Начало'
        hdr_cells[1].text = 'Окончание'
        hdr_cells[2].text = 'Организация'
        hdr_cells[3].text = 'Должность'
        hdr_cells[4].text = 'Обязанности'
        hdr_cells[5].text = 'Причина увольнения'

        for item in work_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("startDate", "—")
            row_cells[1].text = item.get("endDate", "—")
            row_cells[2].text = item.get("organization", "—")
            row_cells[3].text = item.get("position", "—")
            row_cells[4].text = item.get("responsibilities", "—")
            row_cells[5].text = item.get("quitReason", "—")
    else:
        doc.add_paragraph("Опыт работы не указан")

    # Подпись
    doc.add_paragraph("Достоверность указанных сведений: Подтверждено")
    doc.add_paragraph("Согласие на обработку персональных данных: Получено")

    signature_date = form_data['personalInfo'].get("signatureDate", "Не указано")
    p = doc.add_paragraph()
    p.add_run("Дата: ").bold = True
    p.add_run(signature_date)
    p.paragraph_format.space_before = Mm(10)

    # === Сохранение в буфер и конвертация в base64 ===
    buffer = BytesIO()
    doc.save(buffer)  # Сохраняем в буфер
    buffer.seek(0)  # Возвращаем указатель в начало

    base64_content = base64.b64encode(buffer.read()).decode('utf-8')  # Кодируем в base64
    buffer.close()

    # Генерация имени файла
    full_name = form_data['personalInfo'].get("fullName", "соискателя").split()[0]
    filename = f"Анкета_{full_name}_{datetime.now().strftime('%Y-%m-%d')}.docx"
    full_name = form_data['personalInfo'].get("fullName", "Не указано")
    vacancy = form_data['personalInfo'].get("position", "Не указано")

    return base64_content, filename, full_name, vacancy

# Отправка файла в Bitrix24
def send_file(form_data):

    base64_content, filename, full_name, vacancy = save_to_docx(form_data)

    url = "https://imperial44.bitrix24.ru/rest/324/vbmkes7okpbmcmch/disk.folder.uploadfile"

    try:
        # Формируем данные запроса
        payload = {
            "id": "138646",  # ID папки
            "data": {
                "NAME": filename
            },
            "fileContent": base64_content,
            "generateUniqueName": True
        }

        # Отправляем POST-запрос
        response = requests.post(
            url,
            json=payload,  # requests установит Content-Type: application/json
            timeout=30
        )

        # Проверяем статус ответа
        if response.status_code == 200:
            print("Файл создан")
            return response.json()['result']['ID'], full_name, vacancy
        else:
            print(f"Ошибка: {response.status_code}")
            print(f"Текст ответа: {response.text}")
            return None

    except Exception as e:
        print(f"Произошла ошибка: {e}")
        return None

# Создание задачи Bitrix24
def create_task(file_id, full_name):

    url = "https://imperial44.bitrix24.ru/rest/324/vbmkes7okpbmcmch/tasks.task.add"

    today = datetime.now()
    future_date = today + timedelta(days=2)
    formatted_date = future_date.strftime('%Y-%m-%d')

    try:
        # Формируем данные запроса
        payload = {
            "fields": {
                "TITLE": "Анкета " + full_name,
                "DEADLINE": formatted_date,
                "CREATED_BY": 98,
                "RESPONSIBLE_ID": 22,
                "ACCOMPLICES": [104,18],
                "UF_TASK_WEBDAV_FILES": [
                    "n" + str(file_id)
                ],
                "SE_PARAMETER": [
                    {
                        "VALUE": "Y",
                        "CODE": 3
                    }
                ],
                "PARENT_ID": "19024",
                "AUDITORS": [324,176]
            }
        }

        # Отправляем POST-запрос
        response = requests.post(
            url,
            json=payload,  # requests установит Content-Type: application/json
            timeout=30
        )

        # Проверяем статус ответа
        if response.status_code == 200:
            print("Задача поставлена")
            return response.json()['result']['task']['id']
        else:
            print(f"Ошибка: {response.status_code}")
            print(f"Текст ответа: {response.text}")
            return None

    except Exception as e:
        print(f"Произошла ошибка: {e}")
        return None

# Отправка сообщения Bitrix24
def send_message(chat_id, message_text):

    url = "https://imperial44.bitrix24.ru/rest/324/epmzgekyikixd2b9/im.message.add"

    try:
        # Формируем данные запроса
        payload = {
            "DIALOG_ID": chat_id,
			"MESSAGE": message_text
        }

        # Отправляем POST-запрос
        response = requests.post(
            url,
            json=payload,  # requests установит Content-Type: application/json
            timeout=30
        )

        # Проверяем статус ответа
        if response.status_code == 200:
            print("Сообщенеие отправлено")
            return response.json()
        else:
            print(f"Ошибка: {response.status_code}")
            print(f"Текст ответа: {response.text}")
            return None

    except Exception as e:
        print(f"Произошла ошибка: {e}")
        return None

# Проверка задачи по ID Bitrix24
def check_task(task_id):
    url = "https://imperial44.bitrix24.ru/rest/324/vbmkes7okpbmcmch/tasks.task.result.list"

    try:
        # Формируем данные запроса
        payload = {
            "taskId": task_id
        }

        # Отправляем POST-запрос
        response = requests.post(
            url,
            json=payload,  # requests установит Content-Type: application/json
            timeout=30
        )

        # Проверяем статус ответа
        if response.status_code == 200:
            print("Сообщенеие отправлено")
            result = response.json()['result']
            if len(result) == 0:
                return ''
            else:
                return result[0]['text']
        else:
            print(f"Ошибка: {response.status_code}")
            print(f"Текст ответа: {response.text}")
            return None

    except Exception as e:
        print(f"Произошла ошибка: {e}")
        return None
